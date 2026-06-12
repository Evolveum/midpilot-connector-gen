# Copyright (C) 2010-2026 Evolveum and contributors
#
# Licensed under the EUPL-1.2 or later.

import asyncio
import hashlib
import io
import json
import logging
from pathlib import Path
from typing import Any
from uuid import UUID

import yaml
from bs4 import BeautifulSoup
from docx import Document
from fastapi import HTTPException, UploadFile, status
from pypdf import PdfReader

from src.common.chunks import count_tokens, split_text_with_token_overlap
from src.common.database.repositories.session_repository import SessionRepository
from src.common.enums import JobStage
from src.common.jobs import schedule_coroutine_job
from src.common.session.schema import (
    PreparedDocumentationUpload,
    RawUploadedDocumentation,
    SessionUploadContext,
    UploadedDocumentation,
)
from src.config import config

logger = logging.getLogger(__name__)

_CONTENT_TYPES_BY_PARSER = {
    "json": {
        "application/json",
        "application/openapi+json",
        "application/schema+json",
        "application/scim+json",
        "application/sql+json",
        "application/conndev+json",
    },
    "yaml": {
        "application/x-yaml",
        "application/yaml",
        "application/vnd.yaml",
        "application/conndev+yaml",
    },
    "html": {"text/html", "application/xhtml+xml"},
    "text": {
        "application/xml",
        "application/csv",
        "application/sql",
        "application/x-sql",
        "text/sql",
    },
    "pdf": {"application/pdf"},
    "docx": {"application/vnd.openxmlformats-officedocument.wordprocessingml.document"},
}
_SUFFIXES_BY_PARSER = {
    "json": {".json"},
    "yaml": {".yaml", ".yml"},
    "html": {".html", ".htm"},
    "text": {
        ".adoc",
        ".asc",
        ".csv",
        ".graphql",
        ".gql",
        ".log",
        ".md",
        ".openapi",
        ".sql",
        ".txt",
        ".xml",
    },
    "pdf": {".pdf"},
    "docx": {".docx"},
}
_TEXT_SUFFIXES = (
    _SUFFIXES_BY_PARSER["json"]
    | _SUFFIXES_BY_PARSER["yaml"]
    | _SUFFIXES_BY_PARSER["html"]
    | _SUFFIXES_BY_PARSER["text"]
)
_TEXT_CONTENT_TYPES = (
    _CONTENT_TYPES_BY_PARSER["json"]
    | _CONTENT_TYPES_BY_PARSER["yaml"]
    | _CONTENT_TYPES_BY_PARSER["html"]
    | _CONTENT_TYPES_BY_PARSER["text"]
)
_SINGLE_ITEM_SCHEMA_CONTENT_TYPES = {
    "application/scim+json",
    "application/sql+json",
    "application/conndev+json",
    "application/conndev+yaml",
    "application/sql",
    "application/x-sql",
    "text/sql",
}
_GENERIC_CONTENT_TYPES = {"", "application/octet-stream", "binary/octet-stream"}
_CONTENT_TYPE_BY_SUFFIX = {
    ".json": "application/json",
    ".yaml": "application/yaml",
    ".yml": "application/yaml",
    ".html": "text/html",
    ".htm": "text/html",
    ".xml": "application/xml",
    ".sql": "application/sql",
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


def _normalize_content_type(content_type: str | None) -> str:
    raw = (content_type or "").split(";", 1)[0].strip().lower()
    return raw or ""


def _resolve_content_type(content_type: str | None, upload_content_type: str | None, suffix: str) -> str:
    explicit_content_type = _normalize_content_type(content_type)
    if explicit_content_type:
        return explicit_content_type

    normalized_upload_type = _normalize_content_type(upload_content_type)
    if normalized_upload_type not in _GENERIC_CONTENT_TYPES:
        return normalized_upload_type

    return _CONTENT_TYPE_BY_SUFFIX.get(suffix, normalized_upload_type)


def should_preserve_as_single_item(content_type: str | None) -> bool:
    return _normalize_content_type(content_type) in _SINGLE_ITEM_SCHEMA_CONTENT_TYPES


def _decode_text(data: bytes, filename: str) -> str:
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        logger.info("[Upload] Falling back to replacement UTF-8 decode for %s", filename)
        return data.decode("utf-8", errors="replace")


def _pretty_json_or_original(text: str) -> str:
    try:
        return json.dumps(json.loads(text), ensure_ascii=False, indent=2)
    except json.JSONDecodeError:
        return text


def _pretty_yaml_or_original(text: str) -> str:
    try:
        parsed = yaml.safe_load(text)
    except yaml.YAMLError:
        return text
    if parsed is None:
        return text
    return yaml.safe_dump(parsed, allow_unicode=True, sort_keys=False)


def _html_to_text(text: str) -> str:
    soup = BeautifulSoup(text, "html.parser")
    for element in soup(["script", "style"]):
        element.decompose()
    return soup.get_text(separator="\n")


def _pdf_to_text(data: bytes, filename: str) -> tuple[str, dict[str, Any]]:
    try:
        reader = PdfReader(io.BytesIO(data))
        pages: list[str] = []
        for idx, page in enumerate(reader.pages, start=1):
            page_text = page.extract_text() or ""
            if page_text.strip():
                pages.append(f"\n\n--- Page {idx} ---\n{page_text.strip()}")
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Could not extract text from uploaded PDF {filename}.",
        ) from exc

    text = "\n".join(pages).strip()
    if not text:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Uploaded PDF {filename} does not contain extractable text.",
        )
    return text, {"pdf_pages": len(reader.pages)}


def _docx_to_text(data: bytes, filename: str) -> tuple[str, dict[str, Any]]:
    try:
        document = Document(io.BytesIO(data))
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        tables: list[str] = []
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    tables.append(" | ".join(cells))
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Could not extract text from uploaded DOCX {filename}.",
        ) from exc

    text = "\n".join([*paragraphs, *tables]).strip()
    if not text:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Uploaded DOCX {filename} does not contain extractable text.",
        )
    return text, {"docx_paragraphs": len(paragraphs), "docx_tables": len(document.tables)}


def _is_text_upload(content_type: str, suffix: str) -> bool:
    return content_type.startswith("text/") or content_type in _TEXT_CONTENT_TYPES or suffix in _TEXT_SUFFIXES


def _build_metadata(
    *,
    filename: str,
    content_type: str,
    parser: str,
    extra: dict[str, Any] | None = None,
) -> dict[str, Any]:
    metadata = {
        "filename": filename,
        "content_type": content_type,
        "parser": parser,
    }
    if extra:
        metadata.update(extra)
    return metadata


async def read_raw_uploaded_documentation(
    documentation: UploadFile,
    *,
    content_type: str | None = None,
) -> RawUploadedDocumentation:
    filename = documentation.filename or "unknown"
    suffix = Path(filename).suffix.lower()
    content_type = _resolve_content_type(content_type, documentation.content_type, suffix)
    data = await documentation.read()

    if not data:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Uploaded documentation {filename} is empty.",
        )

    return RawUploadedDocumentation(
        data=data,
        filename=filename,
        content_type=content_type,
        content_hash=hashlib.sha256(data).hexdigest(),
    )


async def parse_uploaded_documentation(raw_upload: RawUploadedDocumentation) -> UploadedDocumentation:
    filename = raw_upload.filename
    content_type = raw_upload.content_type
    suffix = Path(filename).suffix.lower()
    data = raw_upload.data
    preserve_as_single_item = should_preserve_as_single_item(content_type)

    parser = "text"
    extra_metadata: dict[str, Any] = {}

    if content_type in _CONTENT_TYPES_BY_PARSER["pdf"] or suffix in _SUFFIXES_BY_PARSER["pdf"]:
        text, extra_metadata = await asyncio.to_thread(_pdf_to_text, data, filename)
        parser = "pdf"
    elif content_type in _CONTENT_TYPES_BY_PARSER["docx"] or suffix in _SUFFIXES_BY_PARSER["docx"]:
        text, extra_metadata = await asyncio.to_thread(_docx_to_text, data, filename)
        parser = "docx"
    elif content_type in _CONTENT_TYPES_BY_PARSER["json"] or suffix in _SUFFIXES_BY_PARSER["json"]:
        text = _pretty_json_or_original(_decode_text(data, filename))
        parser = "json"
    elif content_type in _CONTENT_TYPES_BY_PARSER["yaml"] or suffix in _SUFFIXES_BY_PARSER["yaml"]:
        text = _pretty_yaml_or_original(_decode_text(data, filename))
        parser = "yaml"
    elif content_type in _CONTENT_TYPES_BY_PARSER["html"] or suffix in _SUFFIXES_BY_PARSER["html"]:
        text = _html_to_text(_decode_text(data, filename))
        parser = "html"
    elif _is_text_upload(content_type, suffix):
        text = _decode_text(data, filename)
    else:
        raise HTTPException(
            status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
            detail=(
                f"Unsupported documentation content type '{content_type}' for {filename}. "
                "Supported uploads include JSON, YAML, OpenAPI, Markdown, AsciiDoc, HTML, XML, CSV, SQL, text, PDF, and DOCX."
            ),
        )

    text = text.strip()
    if not text:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Uploaded documentation {filename} did not produce any text.",
        )

    return UploadedDocumentation(
        text=text,
        filename=filename,
        content_type=content_type,
        metadata=_build_metadata(
            filename=filename,
            content_type=content_type,
            parser=parser,
            extra={
                **extra_metadata,
                **(
                    {
                        "preserve_as_single_documentation_item": True,
                        "chunking_strategy": "single_item_schema",
                    }
                    if preserve_as_single_item
                    else {}
                ),
            },
        ),
        preserve_as_single_item=preserve_as_single_item,
    )


async def read_uploaded_documentation(
    documentation: UploadFile,
    *,
    content_type: str | None = None,
) -> UploadedDocumentation:
    raw_upload = await read_raw_uploaded_documentation(documentation, content_type=content_type)
    return await parse_uploaded_documentation(raw_upload)


def chunk_uploaded_documentation(session_id: UUID, uploaded: UploadedDocumentation) -> list[tuple[str, int]]:
    if uploaded.preserve_as_single_item:
        token_count = count_tokens(uploaded.text)
        logger.info(
            "[Upload] Preserving uploaded schema as a single documentation item for session %s "
            "filename=%s content_type=%s tokens=%s",
            session_id,
            uploaded.filename,
            uploaded.content_type,
            token_count,
        )
        return [(uploaded.text, token_count)]

    logger.info(
        "[Upload] Chunking documentation for session %s filename=%s content_type=%s parser=%s",
        session_id,
        uploaded.filename,
        uploaded.content_type,
        uploaded.metadata.get("parser"),
    )
    chunks = split_text_with_token_overlap(
        uploaded.text,
        max_tokens=config.scrape_and_process.chunk_length,
        overlap_ratio=0.05,
    )
    logger.info("[Upload] Generated %s chunks for uploaded document", len(chunks))
    return chunks


async def get_session_upload_context(repo: SessionRepository, session_id: UUID) -> SessionUploadContext:
    session_data = await repo.get_session_data(session_id) or {}
    discovery_input = session_data.get("discoveryInput", {})
    scrape_input = session_data.get("scrapeInput", {})

    return SessionUploadContext(
        app=discovery_input.get("applicationName") or scrape_input.get("applicationName") or "unknown",
        app_version=discovery_input.get("applicationVersion") or scrape_input.get("applicationVersion") or "unknown",
    )


async def prepare_documentation_upload(
    repo: SessionRepository,
    session_id: UUID,
    documentation: UploadFile,
) -> PreparedDocumentationUpload:
    context = await get_session_upload_context(repo, session_id)
    raw_upload = await read_raw_uploaded_documentation(documentation)
    return PreparedDocumentationUpload(raw_upload=raw_upload, context=context)


async def queue_documentation_upload_job(
    *,
    repo: SessionRepository,
    session_id: UUID,
    doc_id: UUID,
    prepared: PreparedDocumentationUpload,
    skip_cache: bool | None = None,
) -> UUID:
    from src.common.session.session import process_documentation_worker

    raw_upload = prepared.raw_upload
    context = prepared.context

    input_payload = {
        "session_id": str(session_id),
        "filename": raw_upload.filename,
        "doc_id": str(doc_id),
        "content_type": raw_upload.content_type,
        "content_hash": raw_upload.content_hash,
        "size_bytes": len(raw_upload.data),
        "app": context.app,
        "app_version": context.app_version,
    }
    if skip_cache is not None:
        input_payload["skipCache"] = skip_cache

    job_id = await schedule_coroutine_job(
        job_type="documentation.processUpload",
        input_payload=input_payload,
        worker=process_documentation_worker,
        worker_kwargs={
            "session_id": session_id,
            "raw_upload": raw_upload,
            "doc_id": doc_id,
            "app": context.app,
            "app_version": context.app_version,
        },
        initial_stage=JobStage.queue,
        initial_message="Queued uploaded documentation for processing",
        session_id=session_id,
    )

    job_key = f"documentation.processUpload_{doc_id}_job_id"
    await repo.update_session(session_id, {job_key: str(job_id)})
    return job_id
