# Copyright (C) 2010-2026 Evolveum and contributors
#
# Licensed under the EUPL-1.2 or later.

import io
import json
import logging
from dataclasses import dataclass
from pathlib import Path
from typing import Any

import yaml
from bs4 import BeautifulSoup
from docx import Document
from fastapi import HTTPException, UploadFile, status
from pypdf import PdfReader

logger = logging.getLogger(__name__)

_TEXT_CONTENT_TYPES = {
    "application/json",
    "application/openapi+json",
    "application/schema+json",
    "application/scim+json",
    "application/x-yaml",
    "application/yaml",
    "application/vnd.yaml",
    "application/xml",
    "application/xhtml+xml",
    "application/csv",
    "application/sql",
    "application/x-sql",
}
_YAML_SUFFIXES = {".yaml", ".yml"}
_JSON_SUFFIXES = {".json"}
_HTML_SUFFIXES = {".html", ".htm"}
_TEXT_SUFFIXES = {
    ".adoc",
    ".asc",
    ".csv",
    ".graphql",
    ".gql",
    ".log",
    ".md",
    ".openapi",
    ".sql",
    ".txt",
    ".xml",
}
_PDF_SUFFIXES = {".pdf"}
_DOCX_SUFFIXES = {".docx"}
_GENERIC_CONTENT_TYPES = {"", "application/octet-stream", "binary/octet-stream"}
_CONTENT_TYPE_BY_SUFFIX = {
    ".json": "application/json",
    ".yaml": "application/yaml",
    ".yml": "application/yaml",
    ".html": "text/html",
    ".htm": "text/html",
    ".xml": "application/xml",
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


@dataclass(frozen=True)
class UploadedDocumentation:
    text: str
    filename: str
    content_type: str
    metadata: dict[str, Any]


def _normalize_content_type(content_type: str | None) -> str:
    raw = (content_type or "").split(";", 1)[0].strip().lower()
    return raw or ""


def _resolve_content_type(content_type: str | None, upload_content_type: str | None, suffix: str) -> str:
    explicit_content_type = _normalize_content_type(content_type)
    if explicit_content_type:
        return explicit_content_type

    normalized_upload_type = _normalize_content_type(upload_content_type)
    if normalized_upload_type not in _GENERIC_CONTENT_TYPES:
        return normalized_upload_type

    return _CONTENT_TYPE_BY_SUFFIX.get(suffix, normalized_upload_type)


def _decode_text(data: bytes, filename: str) -> str:
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        logger.info("[Upload] Falling back to replacement UTF-8 decode for %s", filename)
        return data.decode("utf-8", errors="replace")


def _pretty_json_or_original(text: str) -> str:
    try:
        return json.dumps(json.loads(text), ensure_ascii=False, indent=2)
    except json.JSONDecodeError:
        return text


def _pretty_yaml_or_original(text: str) -> str:
    try:
        parsed = yaml.safe_load(text)
    except yaml.YAMLError:
        return text
    if parsed is None:
        return text
    return yaml.safe_dump(parsed, allow_unicode=True, sort_keys=False)


def _html_to_text(text: str) -> str:
    soup = BeautifulSoup(text, "html.parser")
    for element in soup(["script", "style"]):
        element.decompose()
    return soup.get_text(separator="\n")


def _pdf_to_text(data: bytes, filename: str) -> tuple[str, dict[str, Any]]:
    try:
        reader = PdfReader(io.BytesIO(data))
        pages: list[str] = []
        for idx, page in enumerate(reader.pages, start=1):
            page_text = page.extract_text() or ""
            if page_text.strip():
                pages.append(f"\n\n--- Page {idx} ---\n{page_text.strip()}")
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Could not extract text from uploaded PDF {filename}.",
        ) from exc

    text = "\n".join(pages).strip()
    if not text:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Uploaded PDF {filename} does not contain extractable text.",
        )
    return text, {"pdf_pages": len(reader.pages)}


def _docx_to_text(data: bytes, filename: str) -> tuple[str, dict[str, Any]]:
    try:
        document = Document(io.BytesIO(data))
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        tables: list[str] = []
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    tables.append(" | ".join(cells))
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Could not extract text from uploaded DOCX {filename}.",
        ) from exc

    text = "\n".join([*paragraphs, *tables]).strip()
    if not text:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Uploaded DOCX {filename} does not contain extractable text.",
        )
    return text, {"docx_paragraphs": len(paragraphs), "docx_tables": len(document.tables)}


def _is_text_upload(content_type: str, suffix: str) -> bool:
    return content_type.startswith("text/") or content_type in _TEXT_CONTENT_TYPES or suffix in _TEXT_SUFFIXES


def _build_metadata(
    *,
    filename: str,
    content_type: str,
    original_size: int,
    extracted_text: str,
    parser: str,
    extra: dict[str, Any] | None = None,
) -> dict[str, Any]:
    metadata = {
        "filename": filename,
        "contentType": content_type,
        "original_size": original_size,
        "extracted_length": len(extracted_text),
        "parser": parser,
    }
    if extra:
        metadata.update(extra)
    return metadata


async def read_uploaded_documentation(
    documentation: UploadFile,
    *,
    content_type: str | None = None,
) -> UploadedDocumentation:
    filename = documentation.filename or "unknown"
    suffix = Path(filename).suffix.lower()
    content_type = _resolve_content_type(content_type, documentation.content_type, suffix)
    data = await documentation.read()

    if not data:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Uploaded documentation {filename} is empty.",
        )

    parser = "text"
    extra_metadata: dict[str, Any] = {}

    if content_type == "application/pdf" or suffix in _PDF_SUFFIXES:
        text, extra_metadata = _pdf_to_text(data, filename)
        parser = "pdf"
    elif (
        content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or suffix in _DOCX_SUFFIXES
    ):
        text, extra_metadata = _docx_to_text(data, filename)
        parser = "docx"
    elif (
        content_type
        in {
            "application/json",
            "application/openapi+json",
            "application/schema+json",
            "application/scim+json",
        }
        or suffix in _JSON_SUFFIXES
    ):
        text = _pretty_json_or_original(_decode_text(data, filename))
        parser = "json"
    elif content_type in {"application/x-yaml", "application/yaml", "application/vnd.yaml"} or suffix in _YAML_SUFFIXES:
        text = _pretty_yaml_or_original(_decode_text(data, filename))
        parser = "yaml"
    elif content_type in {"text/html", "application/xhtml+xml"} or suffix in _HTML_SUFFIXES:
        text = _html_to_text(_decode_text(data, filename))
        parser = "html"
    elif _is_text_upload(content_type, suffix):
        text = _decode_text(data, filename)
    else:
        raise HTTPException(
            status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
            detail=(
                f"Unsupported documentation content type '{content_type}' for {filename}. "
                "Supported uploads include JSON, YAML, OpenAPI, Markdown, AsciiDoc, HTML, XML, CSV, SQL, text, PDF, and DOCX."
            ),
        )

    text = text.strip()
    if not text:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Uploaded documentation {filename} did not produce any text.",
        )

    return UploadedDocumentation(
        text=text,
        filename=filename,
        content_type=content_type,
        metadata=_build_metadata(
            filename=filename,
            content_type=content_type,
            original_size=len(data),
            extracted_text=text,
            parser=parser,
            extra=extra_metadata,
        ),
    )
